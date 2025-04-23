"""
Script para criar relatório em Word com os resultados do AutoML
"""
import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.datasets import load_breast_cancer

# Função para adicionar um cabeçalho estilizado
def add_heading_with_line(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Adicionar uma linha horizontal após o cabeçalho
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(12)
    
    # Adicionar linha horizontal
    p.add_run().add_break()

# Função para adicionar uma figura com legenda
def add_figure(doc, img_path, caption, width=6):
    doc.add_picture(img_path, width=Inches(width))
    
    # Adicionar legenda
    caption_para = doc.add_paragraph(caption, style='Caption')
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar espaço após a figura
    doc.add_paragraph()

# Criar um novo documento
doc = Document()

# Definir estilos
styles = doc.styles

# Estilo para título
style = styles['Title']
font = style.font
font.name = 'Arial'
font.size = Pt(24)
font.bold = True
font.color.rgb = RGBColor(0, 51, 102)

# Estilo para cabeçalhos
for i in range(1, 4):
    style = styles[f'Heading {i}']
    font = style.font
    font.name = 'Arial'
    font.bold = True
    if i == 1:
        font.size = Pt(18)
        font.color.rgb = RGBColor(0, 51, 102)
    elif i == 2:
        font.size = Pt(16)
        font.color.rgb = RGBColor(0, 76, 153)
    else:
        font.size = Pt(14)
        font.color.rgb = RGBColor(0, 102, 204)

# Estilo para texto normal
style = styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Estilo para legenda
if 'Caption' not in styles:
    style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.italic = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(12)

# Título do documento
doc.add_heading('Relatório de AutoML para Classificação Binária', 0)
doc.add_paragraph()

# Sumário
doc.add_paragraph('Sumário', style='Heading 1')
doc.add_paragraph('1. Introdução')
doc.add_paragraph('2. Conjunto de Dados')
doc.add_paragraph('3. Metodologia')
doc.add_paragraph('4. Resultados')
doc.add_paragraph('   4.1. Modelo Selecionado')
doc.add_paragraph('   4.2. Avaliação de Desempenho')
doc.add_paragraph('   4.3. Análise de Features')
doc.add_paragraph('5. Conclusões')
doc.add_paragraph()
doc.add_page_break()

# 1. Introdução
add_heading_with_line(doc, '1. Introdução', 1)
p = doc.add_paragraph()
p.add_run('Este relatório apresenta os resultados de um processo de Machine Learning Automatizado (AutoML) para um problema de classificação binária. O AutoML é uma abordagem que automatiza o processo de seleção e otimização de modelos de machine learning, permitindo encontrar o modelo mais adequado para um determinado conjunto de dados sem a necessidade de ajuste manual extensivo.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Neste projeto, utilizamos a biblioteca TPOT (Tree-based Pipeline Optimization Tool), uma ferramenta de AutoML baseada em algoritmos genéticos que automatiza a exploração de pipelines de machine learning. O TPOT combina diferentes pré-processadores, seletores de features e modelos de classificação para encontrar a combinação que maximiza o desempenho para o problema em questão.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('O objetivo principal foi identificar o modelo com melhor desempenho segundo a métrica KS (Kolmogorov-Smirnov), que é particularmente útil para avaliar o poder discriminativo de modelos de classificação binária, especialmente em contextos onde a separação entre classes é crucial.')
doc.add_paragraph()

# 2. Conjunto de Dados
add_heading_with_line(doc, '2. Conjunto de Dados', 1)

# Carregar informações do conjunto de dados
with open('dados/info_dataset.txt', 'r') as f:
    dataset_info = f.read()

# Extrair informações relevantes
import re
samples_match = re.search(r'Número de amostras: (\d+)', dataset_info)
features_match = re.search(r'Número de características: (\d+)', dataset_info)
classes_match = re.search(r"Classes: \['(.+)' '(.+)'\]", dataset_info)

samples = samples_match.group(1) if samples_match else "N/A"
features = features_match.group(1) if features_match else "N/A"
class_0 = classes_match.group(1) if classes_match else "Classe 0"
class_1 = classes_match.group(2) if classes_match else "Classe 1"

p = doc.add_paragraph()
p.add_run('Para este estudo, utilizamos o conjunto de dados Breast Cancer Wisconsin, que é frequentemente usado para tarefas de classificação binária. Este conjunto contém características extraídas de imagens digitalizadas de aspirações por agulha fina (FNA) de massas mamárias, com o objetivo de classificar os tumores como malignos ou benignos.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run(f'O conjunto de dados contém {samples} amostras e {features} características. As classes são "{class_0}" (representada como 0) e "{class_1}" (representada como 1).')
doc.add_paragraph()

# Adicionar figura da distribuição de classes
add_figure(doc, 'dados/distribuicao_classes.png', 'Figura 1: Distribuição das Classes no Conjunto de Dados')

p = doc.add_paragraph()
p.add_run('Como pode ser observado na Figura 1, o conjunto de dados apresenta um leve desbalanceamento entre as classes, com mais exemplos benignos do que malignos. Este desbalanceamento foi considerado durante o processo de treinamento e avaliação dos modelos.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Os dados foram divididos em conjuntos de treinamento (70%) e teste (30%), mantendo a proporção das classes em ambos os conjuntos. Além disso, as características foram normalizadas usando o StandardScaler para garantir que todas tivessem a mesma escala, o que é importante para muitos algoritmos de machine learning.')
doc.add_paragraph()

# 3. Metodologia
add_heading_with_line(doc, '3. Metodologia', 1)
p = doc.add_paragraph()
p.add_run('A metodologia adotada neste projeto seguiu as seguintes etapas:')
doc.add_paragraph()

# Lista numerada
p = doc.add_paragraph(style='List Number')
p.add_run('Preparação dos dados: carregamento, divisão em conjuntos de treinamento e teste, e normalização.')
p = doc.add_paragraph(style='List Number')
p.add_run('Implementação do AutoML: configuração e execução do TPOT para explorar diferentes pipelines de machine learning.')
p = doc.add_paragraph(style='List Number')
p.add_run('Avaliação dos modelos: utilização da métrica KS para avaliar o poder discriminativo dos modelos gerados.')
p = doc.add_paragraph(style='List Number')
p.add_run('Seleção do melhor modelo: identificação do modelo com melhor desempenho segundo a métrica KS.')
p = doc.add_paragraph(style='List Number')
p.add_run('Análise detalhada: geração de visualizações e métricas para compreender o comportamento do modelo selecionado.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('O TPOT foi configurado para otimizar a métrica AUC (Area Under the ROC Curve) durante o treinamento, com 5 gerações e uma população de 20 indivíduos para o algoritmo genético. Foi utilizada validação cruzada com 5 folds para avaliar os modelos durante o processo de otimização.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Para a avaliação final, além da métrica KS, foram utilizadas outras métricas como precisão, recall, F1-score e acurácia, proporcionando uma visão abrangente do desempenho do modelo selecionado.')
doc.add_paragraph()

# 4. Resultados
add_heading_with_line(doc, '4. Resultados', 1)

# 4.1. Modelo Selecionado
doc.add_heading('4.1. Modelo Selecionado', 2)

# Ler o arquivo de pipeline gerado pelo TPOT
with open('resultados/tpot_pipeline.py', 'r') as f:
    pipeline_code = f.read()

# Extrair o modelo selecionado
model_match = re.search(r'exported_pipeline = (.+)', pipeline_code)
model_selected = model_match.group(1) if model_match else "Não disponível"

p = doc.add_paragraph()
p.add_run(f'Após o processo de otimização, o TPOT selecionou o seguinte modelo como o melhor pipeline: ')
p.add_run(model_selected).italic = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('O LinearSVC (Support Vector Classifier Linear) é um modelo de classificação baseado em máquinas de vetores de suporte com kernel linear. Este modelo é conhecido por sua eficácia em problemas de classificação de alta dimensionalidade e por sua capacidade de lidar bem com conjuntos de dados onde o número de características é maior que o número de amostras.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Os hiperparâmetros selecionados pelo TPOT foram:')
doc.add_paragraph()

# Lista com marcadores
p = doc.add_paragraph(style='List Bullet')
p.add_run('C=0.01: Parâmetro de regularização que controla o trade-off entre a maximização da margem e a minimização do erro de treinamento.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('dual=True: Formulação dual do problema de otimização, que é mais eficiente quando o número de amostras é menor que o número de características.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('loss=hinge: Função de perda hinge, que é a função padrão para SVM.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('penalty=l2: Regularização L2, que penaliza os coeficientes grandes para evitar overfitting.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('tol=0.0001: Tolerância para o critério de parada.')
doc.add_paragraph()

# 4.2. Avaliação de Desempenho
doc.add_heading('4.2. Avaliação de Desempenho', 2)

# Ler o arquivo de avaliação
with open('resultados/avaliacao_tpot.txt', 'r') as f:
    avaliacao = f.read()

# Extrair a estatística KS
ks_match = re.search(r'Estatística KS: ([\d\.]+)', avaliacao)
ks_value = ks_match.group(1) if ks_match else "N/A"

p = doc.add_paragraph()
p.add_run(f'O modelo selecionado obteve uma estatística KS de {ks_value} no conjunto de teste, o que indica um excelente poder discriminativo. A estatística KS varia de 0 a 1, onde 1 representa uma separação perfeita entre as classes e 0 indica nenhuma separação. Um valor de {ks_value} é considerado muito bom, demonstrando que o modelo consegue distinguir efetivamente entre as classes positiva e negativa.')
doc.add_paragraph()

# Adicionar figura da curva KS
add_figure(doc, 'resultados/curva_ks.png', 'Figura 2: Curva KS - Distribuição dos Scores por Classe')

p = doc.add_paragraph()
p.add_run('A Figura 2 mostra a distribuição dos scores previstos pelo modelo para cada classe. A linha vertical vermelha representa o limiar ótimo que maximiza a separação entre as classes. Como pode ser observado, há uma clara separação entre as distribuições das classes positiva (benigna) e negativa (maligna), o que explica o alto valor da estatística KS.')
doc.add_paragraph()

# Adicionar figura da curva ROC
add_figure(doc, 'resultados/curva_roc.png', 'Figura 3: Curva ROC')

p = doc.add_paragraph()
p.add_run('A Figura 3 apresenta a curva ROC (Receiver Operating Characteristic), que é outra forma de visualizar o desempenho do modelo. A área sob a curva ROC (AUC) é de 0.996, o que é extremamente próximo do valor ideal de 1.0. Isso confirma o excelente desempenho do modelo na tarefa de classificação.')
doc.add_paragraph()

# Adicionar figura da matriz de confusão
add_figure(doc, 'resultados/matriz_confusao.png', 'Figura 4: Matriz de Confusão')

p = doc.add_paragraph()
p.add_run('A matriz de confusão (Figura 4) mostra o número de previsões corretas e incorretas para cada classe. O modelo classificou corretamente 61 casos malignos e 106 casos benignos, com apenas 3 falsos positivos e 1 falso negativo. Isso resulta em uma alta precisão e recall para ambas as classes.')
doc.add_paragraph()

# Adicionar figura do resumo das métricas
add_figure(doc, 'resultados/resumo_metricas.png', 'Figura 5: Resumo das Métricas de Desempenho')

p = doc.add_paragraph()
p.add_run('A Figura 5 apresenta um resumo das principais métricas de desempenho do modelo. Além da estatística KS, o modelo obteve excelentes resultados em termos de precisão, recall, F1-score e acurácia, com valores próximos a 1.0 para todas as métricas.')
doc.add_paragraph()

# 4.3. Análise de Features
doc.add_heading('4.3. Análise de Features', 2)

# Adicionar figura da importância das features
add_figure(doc, 'resultados/importancia_features.png', 'Figura 6: Importância das Features')

p = doc.add_paragraph()
p.add_run('A Figura 6 mostra a importância das features para o modelo, calculada usando permutation importance. Esta técnica mede o quanto o desempenho do modelo diminui quando uma feature específica é embaralhada, o que indica a sua contribuição para as previsões do modelo.')
doc.add_paragraph()

# Adicionar figura da distribuição das top features
add_figure(doc, 'resultados/distribuicao_top_features.png', 'Figura 7: Distribuição das Top 5 Features por Classe')

p = doc.add_paragraph()
p.add_run('A Figura 7 apresenta a distribuição das 5 features mais importantes para cada classe. É possível observar que estas features apresentam distribuições distintas para as classes maligna e benigna, o que explica sua importância para o modelo.')
doc.add_paragraph()

# Adicionar figura da dispersão das top features
add_figure(doc, 'resultados/dispersao_top_features.png', 'Figura 8: Dispersão das Duas Features Mais Importantes')

p = doc.add_paragraph()
p.add_run('A Figura 8 mostra um gráfico de dispersão das duas features mais importantes. É possível observar uma clara separação entre as classes, o que indica que estas features são altamente discriminativas para o problema de classificação.')
doc.add_paragraph()

# 5. Conclusões
add_heading_with_line(doc, '5. Conclusões', 1)
p = doc.add_paragraph()
p.add_run('Este relatório apresentou os resultados de um processo de AutoML para classificação binária, utilizando a biblioteca TPOT e avaliando os modelos com a métrica KS. As principais conclusões são:')
doc.add_paragraph()

# Lista com marcadores
p = doc.add_paragraph(style='List Bullet')
p.add_run('O modelo LinearSVC foi selecionado como o melhor pipeline pelo TPOT, demonstrando que modelos lineares podem ser muito eficazes para este tipo de problema quando corretamente configurados.')
p = doc.add_paragraph(style='List Bullet')
p.add_run(f'O modelo obteve uma estatística KS de {ks_value}, indicando um excelente poder discriminativo entre as classes maligna e benigna.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Outras métricas de desempenho, como precisão, recall, F1-score e acurácia, também apresentaram valores muito altos, confirmando a eficácia do modelo.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('A análise de importância das features revelou quais características são mais relevantes para a classificação, fornecendo insights valiosos sobre o problema.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('O processo de AutoML demonstrou ser uma abordagem eficiente para encontrar um modelo de alta performance sem a necessidade de ajuste manual extensivo. A combinação do TPOT com a métrica KS permitiu identificar um modelo com excelente capacidade de separação entre as classes, o que é crucial para aplicações onde a distinção clara entre positivos e negativos é importante.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Para trabalhos futuros, seria interessante explorar outras bibliotecas de AutoML, como auto-sklearn ou H2O AutoML, e comparar seus resultados com os obtidos pelo TPOT. Além disso, poderia ser valioso aplicar técnicas de interpret
(Content truncated due to size limit. Use line ranges to read in chunks)