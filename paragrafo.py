from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
import re

def add_markdown_paragraph(document, text):
    """
    Adiciona parágrafo com formatação estilo markdown:
    - **negrito** → texto em negrito
    - *itálico* ou _itálico_ → texto em itálico
    - [link](url) → hiperlink clicável
    
    Args:
        document: Objeto Document
        text (str): Texto com marcações
    """
    paragraph = document.add_paragraph()
    
    # Padrões regex para formatação
    patterns = [
        ('bold', re.compile(r'\*\*(.*?)\*\*')),        # **negrito**
        ('italic', re.compile(r'\*(.*?)\*')),           # *itálico*
        ('italic', re.compile(r'_(.*?)_')),             # _itálico_
        ('link', re.compile(r'\[(.*?)\]\((.*?)\)')),    # [link](url)
    ]
    
    remaining_text = text
    
    while remaining_text:
        # Encontra todas as próximas ocorrências de padrões
        matches = []
        for name, pattern in patterns:
            match = pattern.search(remaining_text)
            if match:
                matches.append((name, match))
        
        if not matches:
            paragraph.add_run(remaining_text)
            break
        
        # Pega a primeira ocorrência mais à esquerda
        first_match = min(matches, key=lambda x: x[1].start())
        name, match = first_match
        
        # Adiciona texto normal antes do padrão
        if match.start() > 0:
            paragraph.add_run(remaining_text[:match.start()])
        
        # Processa o padrão encontrado
        if name == 'bold':
            run = paragraph.add_run(match.group(1))
            run.bold = True
        elif name == 'italic':
            run = paragraph.add_run(match.group(1))
            run.italic = True
        elif name == 'link':
            link_text = match.group(1)
            link_url = match.group(2)
            
            run = paragraph.add_run()
            run.text = link_text
            run.font.color.rgb = RGBColor(0, 0, 255)  # Azul
            run.font.underline = True
            
            # Cria o elemento de hiperlink
            r_id = paragraph.part.relate_to(
                link_url, 
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True
            )
            
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            new_run.append(run._r)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        
        # Atualiza o texto restante
        remaining_text = remaining_text[match.end():]

# Exemplo de uso
doc = Document()

texto = """
Este é um *texto* com **diversas** formatações:
1. **Negrito** para ênfase
2. _Itálico_ para termos técnicos
3. Links como [Google](https://google.com)
4. Combinações como [_**link formatado**_](https://exemplo.com)
"""

add_markdown_paragraph(doc, texto)

# Outro exemplo com múltiplos parágrafos
add_markdown_paragraph(doc, "Segundo *parágrafo* com **outras** formatações.")

doc.save("documento_completo.docx")
